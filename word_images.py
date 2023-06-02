""" @dann_uc """

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, Cm, Mm
# import html2text
from htmldocx import HtmlToDocx


"""------------------------------------- Variables de configuración de párrafos -------------------------------------"""

i = 0                                           # Identacion izquierda del párrafo
fli = 0                                         # Identacion primera linea (sangria francesa para valores negativos)
sb = 0                                          # Espaciado antes del párrafo
sa = 6                                          # Espaciado después del párrafo
ls = 1.15                                       # Interlineado
a = 3                                           # Alineación - 0: Izq, 1: Centro, 2: Derecha, 3: Justificar
fn = 'Cambria'                                  # Establece el tipo de letra
fs = 11                                         # Establece el tamaño de letra
name_style = 'Indent'                           # Nombre del estilo aplicado

"""------------------------------------------- Funciones estáticas --------------------------------------------------"""


def paragraph_style(document, i, fli, sb, sa, ls, a, fn, fs, name_style):
    """
    Formato para los párrafos del documento Resumen Ejecutivo
    :param document:
    :param i:
    :param fli:
    :param sb:
    :param sa:
    :param ls:
    :param a:
    :param fn:
    :param fs:
    :param name_style:
    :return:
    """
    style = document.styles.add_style(name_style, WD_STYLE_TYPE.PARAGRAPH)      # Acceso al diccionario de estilos
    paragraph_format = style.paragraph_format                                   # Propiedad de ParagraphFormat
    paragraph_format.left_indent = Inches(i)                                    # Identacion izquierda del párrafo
    paragraph_format.first_line_indent = Cm(fli)                                # Identacion primera linea (sangria francesa)
    paragraph_format.space_before = Pt(sb)                                      # Espaciado 12 pts antes del párrafo
    paragraph_format.space_after = Pt(sa)                                       # Espaciado 12 pts después del párrafo
    paragraph_format.widow_control = True                                       # Control de viudas y huérfanos
    paragraph_format.line_spacing = ls                                          # Interlineado a 1.5 líneas
    paragraph_format.alignment = a                                              # 0: Izq, 1: Centro, 2: Derecha, 3: Justificar
    font = style.font                                                           # Definiendo el formato del caracter
    font.name = fn                                                              # Establece el tipo de letra
    font.size = Pt(fs)                                                          # Establece el tamaño de letra


"""--------------------------------------------- Cuerpo del documento -----------------------------------------------"""

paragraph1 = "En el círculo quinto está la laguna Estigia, vigilada por Flegias, hijo de Marte. Enfangados en ella, " \
            "luchando eternamente unos contra otros, a golpes y mordiscos, están los condenados por el pecado de la " \
            "ira, mientras que debajo del agua están los acidiosos, es decir, los perezosos y los que vivieron " \
            "tristes y deprimidos sin motivo. Entre los iracundos estaba Felipe Argenti, florentino llamado así " \
            "porque en una ocasión hizo herrar a su caballo con herraduras de plata. Su familia, los Adimari, fue " \
            "la que se quedó con los bienes de Dante cuando fue exiliado. Los cuatro círculos anteriores castigan " \
            "los pecados de incontinencia, mientras que los círculos siguientes castigan (más severamente) la maldad " \
            "en sentido estricto. La región que comprende a todos ellos se llama Dite (un nombre latino para el " \
            "Hades, o el Infierno), una ciudad que toma el color rojo de las llamas que hay en ella por todas " \
            "partes, y está rodeada de murallas, cuyas puertas guardan los demonios."

paragraph3 = "<p>Como ya sabes, todas las páginas web son documentos escritos en diferentes lenguajes de " \
             "programación como el HTML, que es el más extendido. Por cierto que es difícil ver una página " \
             "escrita en un solo lenguaje  y lo normal es que se incluyan fragmentos de otros lenguajes como " \
             "javascript o php que aportan funcionalidades extra a la página.</p><p>Cada página tiene su propio " \
             "código fuente, o sea, unas líneas de programación que determinan el diseño de la página. Si cuando " \
             "visitas un sitio web ves algo bonito en vez de un montón de texto sin sentido para la mayoría de los " \
             "mortales, es gracias a los navegadores, que son capaces de interpretar ese código fuente y " \
             "ofrecérnoslo en un formato visualmente agradable.</p><p>Pero vamos a lo que nos " \
             "interesa.<strong> ¿Cómo ver el código fuente de una página web? </strong>Es muy sencillo. " \
             "Pero si lo que quieres es hacer tu propia web pero no sabes nada de HTML te recomiendo que no " \
             "te comas la cabeza.</p><p><strong>Zed's</strong> dead baby, <em>Zed's</em> dead.</p>"


document = Document()
new_parser = HtmlToDocx()

# style = document.styles.add_style(name_style, WD_STYLE_TYPE.PARAGRAPH)      # Acceso al diccionario de estilos
# paragraph_format = style.paragraph_format                                   # Propiedad de ParagraphFormat
# paragraph_format.left_indent = Inches(i)                                    # Identacion izquierda del párrafo
# paragraph_format.first_line_indent = Cm(fli)                                # Identacion primera linea (sangria francesa)
# paragraph_format.space_before = Pt(sb)                                      # Espaciado 12 pts antes del párrafo
# paragraph_format.space_after = Pt(sa)                                       # Espaciado 12 pts después del párrafo
# paragraph_format.widow_control = True                                       # Control de viudas y huérfanos
# paragraph_format.line_spacing = ls                                          # Interlineado a 1.5 líneas
# paragraph_format.alignment = a                                              # 0: Izq, 1: Centro, 2: Derecha, 3: Justificar
# font = style.font                                                           # Definiendo el formato del caracter
# font.name = fn                                                              # Establece el tipo de letra
# font.size = Pt(fs)                                                          # Establece el tamaño de letra


paragraph_style(document, i, fli, sb, sa, ls, a, fn, fs, name_style)

parrs1 = document.add_paragraph()
parrs1.style = document.styles['Indent']
parrs1.add_run(paragraph1)


p = new_parser.add_html_to_document(paragraph3, document)
# parrs2 = document.add_paragraph()
# parrs2.style = document.styles['Indent']
# parrs2.add_run(p)


# Guardado del documento
document_path = "texto_prueba.docx"
document.save(document_path)




